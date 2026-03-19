"""
Legacy / standalone Discovery server (port DISCOVERY_PORT, default 8081).

生产环境请只使用 payment_server.py：已包含相同的 /api/discovery/* 与配额逻辑，
且 Cloud Run 仅对外暴露一个 PORT。本文件仅保留作本地对照或实验，勿与 payment_server 同时对外提供流量。
"""
from flask import Flask, request, jsonify
from flask_cors import CORS
import sqlite3
import json
import threading
import uuid
import time
import logging
import os
import smtplib
import re
from io import BytesIO
from html import escape
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from datetime import datetime
import requests # For OpenAI-style API calls if needed
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

DB_FILE = 'discovery_tasks.db'

# --- Configuration (To be provided by user via Env Vars) ---
try:
    from dotenv import load_dotenv
    load_dotenv('.env.discovery')
except ImportError:
    pass

OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY') or os.environ.get('OPENROUTER_API_KEY', '')
OPENAI_BASE_URL = os.environ.get('OPENAI_BASE_URL', 'https://openrouter.ai/api/v1')
DEFAULT_MODEL = os.environ.get('OPENAI_MODEL', 'anthropic/claude-sonnet-4.5')
SMTP_SERVER = os.environ.get('SMTP_SERVER', 'smtp.gmail.com')
SMTP_PORT = int(os.environ.get('SMTP_PORT', 587))
SMTP_USER = os.environ.get('SMTP_USER', '')
SMTP_PASSWORD = os.environ.get('SMTP_PASSWORD', '')
SENDER_EMAIL = os.environ.get('SENDER_EMAIL', SMTP_USER)

def get_db_connection():
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db_connection()
    conn.execute('''
        CREATE TABLE IF NOT EXISTS discovery_tasks (
            task_id TEXT PRIMARY KEY,
            user_name TEXT,
            user_email TEXT,
            industry TEXT,
            form_data TEXT,
            status TEXT DEFAULT 'PENDING',
            report_content TEXT,
            error_message TEXT,
            created_at TEXT,
            updated_at TEXT
        )
    ''')
    conn.commit()
    conn.close()
    logger.info("Database initialized.")

init_db()

# --- AI Analysis & Email Logic ---

def generate_report(form_data):
    """
    Call AI API to generate the discovery report based on form data.
    输出 8000+ 汉字的深度报告，对标付费竞品分析质量。
    """
    if not OPENAI_API_KEY:
        return "ERROR: OpenAI API Key not configured."

    model_to_use = form_data.get('ai_model', DEFAULT_MODEL)
    report_lang = form_data.get('report_language', 'zh')

    lang_instruction = """
    [语言要求]:
    1. 报告主体必须使用简体中文撰写。
    2. 亚马逊关键词、产品术语、用户原话引用保留原始语言（英文/德文/日文等）。
    """ if report_lang == 'zh' else """
    [Language Requirement]:
    1. The report must be written entirely in English.
    2. Keep Amazon keywords and product terms in their original form.
    """

    prompt = f"""你是一位顶级的亚马逊选品与竞品分析专家，曾为数百个卖家撰写过单份售价数千元的深度报告。请基于以下市场发现请求，生成一份**专业级全品类深度分析报告**，结构与深度对标 FlowAI Agent 付费竞品分析报告标准。

【字数硬性要求】报告正文必须达到 8000 汉字以上（英文报告 4000+ words），务必写足写满，不要提前结束。每个表格需有 4-8 行实质内容，每个小节需展开论述（每段 80-150 字），切忌泛泛而谈。

[输入参数]
- 主类目: {form_data.get('category_main')}
- 类目路径: {form_data.get('category_path')}
- 产品关键词: {form_data.get('keywords')}
- 目标站点: {form_data.get('marketplace')}
- 分析重点: {', '.join(form_data.get('focus_areas', []))}
- 目标价格区间: ${form_data.get('target_price_min') or '—'} - ${form_data.get('target_price_max') or '—'}
- 参考ASIN (如有): {', '.join(form_data.get('reference_asins', [])) or '无'}
- 个性化要求: {form_data.get('custom_prompt') or '无'}
{lang_instruction}

[报告结构 - 严格按以下六大部分撰写，每部分需有实质数据与可执行建议]

---

# 全品类深度分析报告: [品类名称]

## 一、市场吸引力综合评估

### 1.1 市场评分卡
用表格呈现，每维度 1-10 分，并附 80-120 字详细分析：
| 维度 | 评分(1-10) | 详细分析 |
| 需求强度 | X | [分析该品类的需求稳定性、复购性、季节性、用户群体规模] |
| 竞争壁垒 | X | [分析品牌垄断度、技术门槛、供应链难度、新进入者空间] |
| 利润潜力 | X | [分析成本结构、毛利率空间、LTV、价格战风险] |
| 进入难度 | X | [分析资金需求、供应链、认证、初期获客难度] |

**综合市场吸引力评分: X.X/10**
**综合判断: ENTER / CAUTIOUS ENTER / NO-GO**
**详细理由:** [200-300字，结合上述维度给出进入或谨慎进入的具体理由与风险提示]

### 1.2 价格生态系统深度分析
- **价格影响因素矩阵** (表格): 因素 | 对价格的影响 | 说明
- **推演的价格带分布**: 引流款区间、主流款区间、高端款区间，每区间需说明特征、竞争程度、策略建议
- **多包装规格价格策略建议** (表格): 规格 | 建议定价 | 单片/单位价格 | 目标用户 | 战略目的
- **建议切入价位**: 给出具体 $XX.XX，并详述 5-6 条定价理由（心理定价、成本结构、竞争定位、促销空间等）

---

## 二、用户群体深度画像

### 2.1 核心用户画像表
| 画像编号 | 用户类型 | 年龄 | 职业特征 | 购买动机 | 价格敏感度 | 占比估计 |
[4-5 行，每类用户有具体描述]

### 2.2 场景化需求分析表
| 优先级 | 使用场景 | 详细描述 | 转化驱动因素 | 评论证据/推理 |
[5-8 行，每个场景 50-80 字描述，转化驱动因素 2-3 条]

### 2.3 用户痛点深度挖掘
| 排名 | 痛点 | 占比 | 根因分析 | 影响用户群 | 解决难度 | 差异化价值 |
[5-8 个痛点，每个有根因分析与对策方向]

### 2.4 用户决策路径分析
- **专业用户/批量采购者** 的购买决策流程 (5-8 步)
- **个人消费者/DIY 用户** 的购买决策流程 (5-8 步)
- **关键决策节点** 与 Listing 优化对应建议

---

## 三、产品与技术趋势

### 3.1 材质/设计演进表
| 阶段 | 主流材质/技术 | 核心功能 | 优势 | 劣势 | 市场状态 |
[过去、当前、未来趋势，3-5 行]

### 3.2 规格/型号趋势分析
针对该品类的关键规格维度（如粒度、尺寸、功率等），分析各规格的市场需求趋势

### 3.3 爆款基因解码
- **必备功能 (Table Stakes)**: 表格列出 5-6 项，不具备就出局
- **加分功能 (Points of Difference)**: 表格列出 4-5 项，具备可脱颖而出
- **设计语言**: 色彩、信息层级、主图要求、描述风格
- **定价区间**: 基于成功产品的价格带分析

### 3.4 技术创新机会
| 创新方向 | 技术描述 | 解决的问题 | 可行性 | 潜在影响 |
[3-5 个创新方向]

---

## 四、竞争格局详解

### 4.1 品牌垄断度分析
- 市场结构评估 (头部集中度、中小品牌空间、白牌比例、价格竞争度)
- 头部品牌识别 (3-5 个典型品牌及定位)
- 市场格局判断: 垄断型/寡头型/分散竞争型
- 新卖家切入可能性: 高/中/低，及理由

### 4.2 竞品深度解剖
选取 1-2 个典型竞品（可基于品类知识虚构或使用"假设竞品A/B"），每个包含：
- 基本信息 (价格、评分、评论数、核心规格)
- 核心优势 (从品类共性推断，3-5 条，每条附"评论证据"风格的说明)
- 核心劣势 (3-4 条，含根因分析)
- 市场定位与竞争策略建议

### 4.3 竞品对比矩阵与市场空白分析
- 维度对比表格 (价格、评分、材料、功能、包装、目标用户等)
- **识别的市场空白和机会** (表格): 机会编号 | 市场空白描述 | 目标用户 | 产品方案 | 预期定价 | 差异化价值
[4-6 个具体机会]
- **推荐的市场进入策略**: 三阶段策略 (0-3月、3-6月、6-12月)

---

## 五、蓝海机会与差异化策略

### 5.1 未被满足的需求深度挖掘
| 排名 | 未满足需求 | 严重程度 | 证据/推理 | 影响用户群 | 解决难度 | 差异化价值 | 实施优先级 |
[5-8 项，每项有实质分析]

### 5.2 差异化策略矩阵
| 策略类型 | 具体方案 | 实施成本 | 实施周期 | 预期效果 | 风险 | 优先级 |
[6-10 项，涵盖产品创新、包装创新、内容创新、服务创新、定位创新、SKU 策略]

### 5.3 推荐的差异化组合策略
- **阶段一 (0-2 月)**: 低成本快速差异化，3-4 项具体措施
- **阶段二 (2-4 月)**: 产品线扩展，2-3 项
- **阶段三 (4-6 月)**: 技术差异化，2-3 项

### 5.4 差异化价值主张设计
- 核心价值主张 (一句话)
- 支撑点 (4-5 条)
- 针对不同用户群的定制化价值主张 (表格)

---

## 六、进入策略与执行计划

### 6.1 理想产品规格定义
- 维度 | 规格要求 | 理由 | 供应商要求 (表格)
- 目标成本结构 (采购、运输、FBA、佣金、广告、目标净利润)
- 目标售价与定价理由

### 6.2 供应链与质量控制计划
- 供应商筛选标准 (6-8 项)
- 供应商开发流程 (4-5 阶段)
- 首批订单建议 (数量、分配、总投资)
- 质量控制流程 (生产前/中/后/到货/上架)

### 6.3 风险雷达与缓解措施
| 风险类型 | 具体风险 | 概率 | 影响 | 缓解措施 | 应急预案 |
[8-12 项，涵盖市场、供应链、运营、合规、财务、平台风险]

### 6.4 30-60-90 天执行计划
- **第 1-30 天 (筹备期)**: 按周分解，每周 3-5 个具体任务，含关键产出
- **第 31-60 天 (生产和启动期)**: 同上
- **第 61-90 天 (验证和优化期)**: 同上
- 每阶段需有明确的成功标准与决策点 (验证成功/需优化/考虑退出)

---

**重要**: 整份报告必须不少于 8000 汉字。基于品类知识进行合理推断，数据可标注为"基于行业经验推断"或"合理估算"。每个表格和列表都要有实质信息，避免空洞概括。"""

    def call_model(messages):
        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://flowaiagent.com",
            "X-Title": "FlowAI Agent"
        }
        payload = {
            "model": model_to_use,
            "messages": messages,
            "temperature": 0.6,
            "max_tokens": 16384
        }

        logger.info(f"Generating report using model: {model_to_use} (target 8000+ chars)")
        response = requests.post(f"{OPENAI_BASE_URL}/chat/completions", json=payload, headers=headers, timeout=300)
        if not response.ok:
            err_body = response.text[:500] if response.text else "(empty)"
            logger.error(f"OpenRouter API Error: status={response.status_code}, body={err_body}")
            return f"ERROR: OpenRouter API failed (HTTP {response.status_code}). Check credits at openrouter.ai/settings/credits. Details: {err_body}"
        data = response.json()
        choice = data['choices'][0]
        content = choice['message']['content']
        finish_reason = choice.get('finish_reason')
        return content, finish_reason

    def report_looks_complete(content):
        normalized = re.sub(r'\s+', '', content)
        completion_markers = [
            '30-60-90天执行计划',
            '第61-90天',
            '6.4',
            '##六、进入策略与执行计划',
            '###6.4',
            '第31-60天',
        ]
        return any(marker in normalized for marker in completion_markers)

    try:
        messages = [{"role": "user", "content": prompt}]
        content, finish_reason = call_model(messages)
        if isinstance(content, str) and content.startswith("ERROR:"):
            return content

        combined = content
        logger.info(f"Report generated chunk 1: {len(content)} chars, finish_reason={finish_reason}")

        for attempt in range(2):
            if finish_reason != 'length' and report_looks_complete(combined):
                break
            continuation_prompt = (
                "Continue the same report from exactly where you stopped. "
                "Do not repeat previous sections. Resume from the next unfinished heading and finish the remaining sections."
            )
            messages = [
                {"role": "user", "content": prompt},
                {"role": "assistant", "content": combined},
                {"role": "user", "content": continuation_prompt}
            ]
            next_chunk, finish_reason = call_model(messages)
            if isinstance(next_chunk, str) and next_chunk.startswith("ERROR:"):
                break
            if not next_chunk.strip():
                break
            combined += "\n\n" + next_chunk.strip()
            logger.info(
                f"Report continuation chunk {attempt + 2}: {len(next_chunk)} chars, "
                f"total={len(combined)}, finish_reason={finish_reason}"
            )

        logger.info(f"Report generated: {len(combined)} chars total")
        return combined
    except requests.exceptions.RequestException as e:
        err_detail = str(e)
        if hasattr(e, 'response') and e.response is not None:
            err_detail = f"HTTP {e.response.status_code}: {e.response.text[:300] if e.response.text else 'no body'}"
        logger.error(f"AI Request Error: {err_detail}")
        return f"ERROR: API request failed. {err_detail}"
    except Exception as e:
        logger.error(f"AI Generation Error: {e}")
        return f"ERROR: Failed to generate report. {str(e)}"

def send_email(to_email, user_name, report_content):
    """
    Send the generated report to the user's email.
    """
    if not SMTP_USER or not SMTP_PASSWORD:
        logger.warning("SMTP not configured. Skipping email.")
        return False

    def render_inline_markdown(text):
        escaped = escape(text)
        escaped = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', escaped)
        escaped = re.sub(r'`([^`]+)`', r'<code>\1</code>', escaped)
        return escaped

    def markdown_to_basic_html(markdown_text):
        parts = []
        list_open = False
        table_rows = []

        def flush_list():
            nonlocal list_open
            if list_open:
                parts.append('</ul>')
                list_open = False

        def flush_table():
            nonlocal table_rows
            if not table_rows:
                return

            def is_separator(row):
                return all(cell and set(cell) <= set('-: ') for cell in row)

            rows = table_rows[:]
            table_rows = []
            if len(rows) >= 2 and is_separator(rows[1]):
                header = rows[0]
                body = rows[2:]
            else:
                header = rows[0]
                body = rows[1:]

            parts.append('<table style="width:100%;border-collapse:collapse;margin:16px 0;">')
            parts.append('<thead><tr>')
            for cell in header:
                parts.append(
                    f'<th style="border:1px solid #dbe3f0;padding:8px 10px;background:#f8fafc;text-align:left;">{render_inline_markdown(cell)}</th>'
                )
            parts.append('</tr></thead>')
            if body:
                parts.append('<tbody>')
                for row in body:
                    parts.append('<tr>')
                    for cell in row:
                        parts.append(
                            f'<td style="border:1px solid #dbe3f0;padding:8px 10px;vertical-align:top;">{render_inline_markdown(cell)}</td>'
                        )
                    parts.append('</tr>')
                parts.append('</tbody>')
            parts.append('</table>')

        for raw_line in markdown_text.splitlines():
            line = raw_line.rstrip()
            stripped = line.strip()

            if stripped.startswith('|') and stripped.endswith('|'):
                flush_list()
                cells = [cell.strip() for cell in stripped.strip('|').split('|')]
                table_rows.append(cells)
                continue
            flush_table()

            if not stripped:
                flush_list()
                continue

            if stripped.startswith('# '):
                flush_list()
                parts.append(f'<h1 style="font-size:24px;margin:28px 0 14px;">{render_inline_markdown(stripped[2:])}</h1>')
            elif stripped.startswith('## '):
                flush_list()
                parts.append(f'<h2 style="font-size:20px;margin:24px 0 12px;">{render_inline_markdown(stripped[3:])}</h2>')
            elif stripped.startswith('### '):
                flush_list()
                parts.append(f'<h3 style="font-size:17px;margin:20px 0 10px;">{render_inline_markdown(stripped[4:])}</h3>')
            elif stripped.startswith('- ') or stripped.startswith('* '):
                if not list_open:
                    parts.append('<ul style="margin:12px 0 12px 22px;padding:0;">')
                    list_open = True
                parts.append(f'<li style="margin:6px 0;">{render_inline_markdown(stripped[2:])}</li>')
            else:
                flush_list()
                parts.append(f'<p style="margin:12px 0;line-height:1.75;">{render_inline_markdown(stripped)}</p>')

        flush_list()
        flush_table()
        return ''.join(parts)

    def build_preview(text, max_lines=18, max_chars=2200):
        lines = []
        total = 0
        for raw_line in text.splitlines():
            stripped = raw_line.strip()
            if not stripped or stripped.startswith('|'):
                continue
            if set(stripped) <= set('-:| '):
                continue
            cleaned = re.sub(r'^[#*\-\d\.\s]+', '', stripped).strip()
            if not cleaned:
                continue
            lines.append(cleaned)
            total += len(cleaned)
            if len(lines) >= max_lines or total >= max_chars:
                break
        return '\n'.join(lines)

    def strip_markdown_markers(text):
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        return text.strip()

    def report_to_docx_bytes(text, title_text):
        document = Document()
        normal_style = document.styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        normal_style.font.size = Pt(10.5)

        table_rows = []

        def flush_table():
            nonlocal table_rows
            if not table_rows:
                return

            def is_separator(row):
                return all(cell and set(cell) <= set('-: ') for cell in row)

            rows = table_rows[:]
            table_rows = []
            if len(rows) >= 2 and is_separator(rows[1]):
                header = rows[0]
                body = rows[2:]
            else:
                header = rows[0]
                body = rows[1:]

            table = document.add_table(rows=1 + len(body), cols=len(header))
            table.style = 'Table Grid'
            for idx, cell in enumerate(header):
                table.rows[0].cells[idx].text = strip_markdown_markers(cell)
            for row_idx, row in enumerate(body, start=1):
                for col_idx, cell in enumerate(row):
                    table.rows[row_idx].cells[col_idx].text = strip_markdown_markers(cell)
            document.add_paragraph('')

        if title_text:
            document.add_heading(strip_markdown_markers(title_text), level=0)

        for raw_line in text.splitlines():
            stripped = raw_line.strip()
            if stripped.startswith('|') and stripped.endswith('|'):
                cells = [cell.strip() for cell in stripped.strip('|').split('|')]
                table_rows.append(cells)
                continue
            flush_table()
            if not stripped:
                continue
            if stripped.startswith('# '):
                document.add_heading(strip_markdown_markers(stripped[2:]), level=1)
            elif stripped.startswith('## '):
                document.add_heading(strip_markdown_markers(stripped[3:]), level=2)
            elif stripped.startswith('### '):
                document.add_heading(strip_markdown_markers(stripped[4:]), level=3)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                document.add_paragraph(strip_markdown_markers(stripped[2:]), style='List Bullet')
            else:
                document.add_paragraph(strip_markdown_markers(stripped))

        flush_table()

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    msg = MIMEMultipart('mixed')
    msg['From'] = SENDER_EMAIL
    msg['To'] = to_email
    msg['Subject'] = f"Your Product Discovery Report for {user_name}"

    safe_name = re.sub(r'[^A-Za-z0-9_-]+', '_', (user_name or 'user')).strip('_') or 'user'
    preview_text = build_preview(report_content)
    report_docx = report_to_docx_bytes(report_content, "Product Discovery Report")

    plain_body = (
        f"Hello {user_name},\n\n"
        "Your Amazon Product Discovery report is ready.\n\n"
        "For better readability and to avoid email clipping, the full report is attached as a DOCX document.\n\n"
        "Preview:\n"
        "----------------------------------------\n"
        f"{preview_text}\n\n"
        "Attachments:\n"
        f"- product_discovery_report_{safe_name}.docx\n"
    )

    html_body = f"""
    <html>
      <body style="margin:0;padding:24px;background:#f3f6fb;color:#0f172a;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;">
        <div style="max-width:840px;margin:0 auto;background:#ffffff;border:1px solid #e2e8f0;border-radius:16px;padding:32px;">
          <h1 style="margin:0 0 12px;font-size:28px;">Product Discovery Report</h1>
          <p style="margin:0 0 12px;line-height:1.7;">Hi {escape(user_name)}, your report is ready.</p>
          <p style="margin:0 0 18px;line-height:1.7;">
            To avoid long-email clipping, the <strong>full report is attached</strong> as a DOCX document that opens cleanly in Word / Google Docs.
            The email body only shows a readable preview.
          </p>
          <div style="background:#f8fafc;border:1px solid #dbe3f0;border-radius:12px;padding:18px 20px;margin:20px 0;">
            {markdown_to_basic_html(preview_text)}
          </div>
          <p style="margin:18px 0 0;color:#475569;">Attachments:</p>
          <ul style="margin:8px 0 0 20px;line-height:1.8;color:#475569;">
            <li>product_discovery_report_{safe_name}.docx</li>
          </ul>
        </div>
      </body>
    </html>
    """

    alternative = MIMEMultipart('alternative')
    alternative.attach(MIMEText(plain_body, 'plain', 'utf-8'))
    alternative.attach(MIMEText(html_body, 'html', 'utf-8'))
    msg.attach(alternative)

    docx_attachment = MIMEApplication(
        report_docx,
        _subtype='vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    docx_attachment.add_header('Content-Disposition', 'attachment', filename=f'product_discovery_report_{safe_name}.docx')
    msg.attach(docx_attachment)

    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SMTP_USER, SMTP_PASSWORD)
        server.send_message(msg)
        server.quit()
        logger.info(f"Email sent to {to_email}")
        return True
    except Exception as e:
        logger.error(f"Email Error: {e}")
        return False

def background_worker(task_id, user_name, user_email, form_data):
    """
    Background process to generate report and send email.
    """
    try:
        # Update status to ANALYZING
        conn = get_db_connection()
        conn.execute("UPDATE discovery_tasks SET status = ?, updated_at = ? WHERE task_id = ?", 
                     ('ANALYZING', datetime.now().isoformat(), task_id))
        conn.commit()

        # 1. Generate AI Report
        report = generate_report(form_data)
        
        if report.startswith("ERROR:"):
            raise Exception(report)

        # 2. Send Email
        email_sent = send_email(user_email, user_name, report)

        # 3. Final Update
        conn.execute("UPDATE discovery_tasks SET status = ?, report_content = ?, updated_at = ? WHERE task_id = ?", 
                     ('COMPLETED' if email_sent else 'COMPLETED_NO_EMAIL', report, datetime.now().isoformat(), task_id))
        conn.commit()
        conn.close()
        logger.info(f"Task {task_id} completed successfully.")

    except Exception as e:
        logger.error(f"Worker Error for task {task_id}: {e}")
        conn = get_db_connection()
        conn.execute("UPDATE discovery_tasks SET status = ?, error_message = ?, updated_at = ? WHERE task_id = ?", 
                     ('FAILED', str(e), datetime.now().isoformat(), task_id))
        conn.commit()
        conn.close()

# --- API Endpoints ---

@app.route('/api/discovery/submit', methods=['POST'])
def submit_task():
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400

    task_id = str(uuid.uuid4())
    user_name = data.get('user_name', 'User')
    user_email = data.get('user_email', '')
    industry = data.get('industry', '')
    
    if not user_email:
        return jsonify({"error": "Email is required"}), 400

    # Store task in DB
    conn = get_db_connection()
    conn.execute('''
        INSERT INTO discovery_tasks (task_id, user_name, user_email, industry, form_data, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (task_id, user_name, user_email, industry, json.dumps(data), 
          datetime.now().isoformat(), datetime.now().isoformat()))
    conn.commit()
    conn.close()

    # Start background worker
    thread = threading.Thread(target=background_worker, args=(task_id, user_name, user_email, data))
    thread.start()

    logger.info(f"Submitted discovery task {task_id} for {user_email}")
    return jsonify({"task_id": task_id, "status": "PENDING"}), 202

@app.route('/api/discovery/status', methods=['GET'])
def get_status():
    task_id = request.args.get('task_id')
    if not task_id:
        return jsonify({"error": "task_id required"}), 400

    conn = get_db_connection()
    row = conn.execute("SELECT status, error_message, updated_at FROM discovery_tasks WHERE task_id = ?", (task_id,)).fetchone()
    conn.close()

    if not row:
        return jsonify({"error": "Task not found"}), 404

    return jsonify(dict(row))

if __name__ == '__main__':
    port = int(os.environ.get('DISCOVERY_PORT', 8081))
    logger.info(f"Discovery Server running on port {port}")
    app.run(host='0.0.0.0', port=port)
