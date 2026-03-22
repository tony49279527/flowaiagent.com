"""
Unified Payment Polling Server (SQLite Version)

Features:
1. Stores payment status in SQLite database (orders.db)
2. Provides /api/check_status for frontend polling
3. Provides /api/update_status for admin trigger
4. Auto-creates SQLite table on startup

Usage:
1. Run: python payment_server.py
2. Expose: ngrok http 8080
"""

from flask import Flask, jsonify, request, send_from_directory, redirect
from flask_cors import CORS
import logging
from datetime import datetime, timedelta
import os
import requests
import threading
import sqlite3
import json
import random
import time
import uuid
import smtplib
import re
from io import BytesIO
from html import escape
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
try:
    from google.cloud import firestore
except ImportError:
    firestore = None

# Serve static files from current directory
app = Flask(__name__)

try:
    from dotenv import load_dotenv
    load_dotenv('.env.discovery')
    load_dotenv('.env') # Also load main .env if it exists
except ImportError:
    pass

ADMIN_API_TOKEN = os.environ.get('ADMIN_API_TOKEN')

# --- Product Discovery (选品分析) Config ---
# 支持 OPENAI_API_KEY 或 OPENROUTER_API_KEY（Cloud Run 任选其一即可）
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY') or os.environ.get('OPENROUTER_API_KEY', '')
OPENAI_BASE_URL = os.environ.get('OPENAI_BASE_URL', 'https://openrouter.ai/api/v1')
OPENAI_MODEL = os.environ.get('OPENAI_MODEL', 'anthropic/claude-sonnet-4.5')
SMTP_SERVER = os.environ.get('SMTP_SERVER', 'smtp.gmail.com')
SMTP_PORT = int(os.environ.get('SMTP_PORT', 587))
SMTP_USER = os.environ.get('SMTP_USER', '')
SMTP_PASSWORD = os.environ.get('SMTP_PASSWORD', '')
SENDER_EMAIL = os.environ.get('SENDER_EMAIL', SMTP_USER)
DISCOVERY_DB_FILE = 'discovery_tasks.db'
cors_origins = [
    origin.strip()
    for origin in os.environ.get(
        'CORS_ORIGINS',
        'https://flowaiagent.com,http://localhost:4173,http://127.0.0.1:4173'
    ).split(',')
    if origin.strip()
]
if cors_origins:
    CORS(app, resources={r"/api/*": {"origins": cors_origins}})
else:
    CORS(app)

# Configure Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

DB_FILE = 'orders.db'
FREE_QUOTA_LIMIT = 2
VALID_QUOTA_FEATURES = {'competitor', 'discovery'}
PERSISTENCE_BACKEND = os.environ.get('PERSISTENCE_BACKEND', 'auto').strip().lower()
FIRESTORE_QUOTA_COLLECTION = os.environ.get('FIRESTORE_QUOTA_COLLECTION', 'user_quota')
FIRESTORE_DISCOVERY_COLLECTION = os.environ.get('FIRESTORE_DISCOVERY_COLLECTION', 'discovery_tasks')

LOCALHOSTS = {'127.0.0.1', '::1'}
_firestore_client = None

def is_admin_request():
    if ADMIN_API_TOKEN:
        token = request.headers.get('X-Admin-Token') or request.args.get('admin_token')
        return token == ADMIN_API_TOKEN
    return request.remote_addr in LOCALHOSTS

def get_db_connection():
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    return conn

def normalize_email(email):
    return str(email or '').strip().lower()

def normalize_quota_feature(feature):
    raw = str(feature or '').strip().lower()
    if raw in ('discovery', 'product_discovery'):
        return 'discovery'
    return 'competitor'

def get_firestore_client():
    global _firestore_client
    if PERSISTENCE_BACKEND == 'sqlite' or firestore is None:
        return None
    if _firestore_client is None:
        try:
            client = firestore.Client()
            # Probe the configured database once so we can safely fall back
            # to SQLite if the API is enabled but the default database does
            # not exist yet or the runtime service account lacks permissions.
            list(client.collection(FIRESTORE_QUOTA_COLLECTION).limit(1).stream())
            _firestore_client = client
        except Exception as e:
            logger.warning(f"Firestore unavailable, falling back to SQLite: {e}")
            _firestore_client = False
    return _firestore_client or None

def using_firestore():
    return get_firestore_client() is not None

def quota_doc_ref(email):
    client = get_firestore_client()
    if client is None:
        return None
    return client.collection(FIRESTORE_QUOTA_COLLECTION).document(normalize_email(email))

def discovery_doc_ref(task_id):
    client = get_firestore_client()
    if client is None:
        return None
    return client.collection(FIRESTORE_DISCOVERY_COLLECTION).document(str(task_id))

def parse_order_id(raw_order_id):
    if raw_order_id is None:
        return None
    try:
        return int(str(raw_order_id).strip())
    except (TypeError, ValueError):
        return None


def normalize_amazon_asin_list(raw):
    """Normalize ASINs from list or string; dedupe; 10-char alphanumeric starting with B0 (aligned with frontend)."""
    items = []
    if raw is None:
        return items
    if isinstance(raw, str):
        chunks = re.split(r'[\s,;]+', raw.strip())
        items = [c.strip().upper() for c in chunks if c and c.strip()]
    elif isinstance(raw, (list, tuple)):
        for x in raw:
            if x is None:
                continue
            s = str(x).strip()
            if not s:
                continue
            chunks = re.split(r'[\s,;]+', s)
            items.extend(c.strip().upper() for c in chunks if c.strip())
    seen = set()
    out = []
    for p in items:
        clean = re.sub(r'[^A-Z0-9]', '', p)
        if len(clean) == 10 and clean.startswith('B0') and clean not in seen:
            seen.add(clean)
            out.append(clean)
    return out

def init_db():
    try:
        conn = get_db_connection()
        conn.execute('''
            CREATE TABLE IF NOT EXISTS orders (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                status TEXT NOT NULL DEFAULT 'PENDING',
                order_data TEXT,
                updated_at TEXT,
                email TEXT,
                quota_usage INTEGER DEFAULT 0
            )
        ''')
        conn.execute('''
            CREATE TABLE IF NOT EXISTS user_quota (
                email TEXT PRIMARY KEY,
                competitor_usage INTEGER NOT NULL DEFAULT 0,
                discovery_usage INTEGER NOT NULL DEFAULT 0,
                updated_at TEXT
            )
        ''')
        # Initialize default row if not exists (for single-user demo)
        cur = conn.execute('SELECT * FROM orders WHERE id = 1')
        if not cur.fetchone():
            conn.execute("INSERT INTO orders (id, status, updated_at) VALUES (1, 'PENDING', ?)", (datetime.now().isoformat(),))
            logger.info("Initialized default order row.")
        conn.commit()
        conn.close()
    except Exception as e:
        logger.error(f"DB Init Error: {e}")

# Initialize DB on start
init_db()

# --- Product Discovery DB ---
def get_discovery_db():
    conn = sqlite3.connect(DISCOVERY_DB_FILE)
    conn.row_factory = sqlite3.Row
    return conn

def init_discovery_db():
    try:
        conn = get_discovery_db()
        conn.execute('''
            CREATE TABLE IF NOT EXISTS discovery_tasks (
                task_id TEXT PRIMARY KEY,
                user_name TEXT,
                user_email TEXT,
                industry TEXT,
                form_data TEXT,
                status TEXT DEFAULT 'PENDING',
                report_content TEXT,
                error_message TEXT,
                created_at TEXT,
                updated_at TEXT
            )
        ''')
        conn.commit()
        conn.close()
        logger.info("Discovery DB initialized.")
    except Exception as e:
        logger.error(f"Discovery DB Init Error: {e}")

init_discovery_db()

def create_discovery_task_record(task_id, user_name, user_email, industry, form_data):
    now = datetime.now().isoformat()
    payload = {
        'task_id': str(task_id),
        'user_name': user_name,
        'user_email': normalize_email(user_email),
        'industry': industry,
        'form_data': json.dumps(form_data),
        'status': 'PENDING',
        'report_content': '',
        'error_message': '',
        'created_at': now,
        'updated_at': now,
    }
    if using_firestore():
        discovery_doc_ref(task_id).set(payload)
        return payload

    conn = get_discovery_db()
    conn.execute('''
        INSERT INTO discovery_tasks (task_id, user_name, user_email, industry, form_data, status, report_content, error_message, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        payload['task_id'],
        payload['user_name'],
        payload['user_email'],
        payload['industry'],
        payload['form_data'],
        payload['status'],
        payload['report_content'],
        payload['error_message'],
        payload['created_at'],
        payload['updated_at'],
    ))
    conn.commit()
    conn.close()
    return payload

def update_discovery_task_record(task_id, **updates):
    update_payload = dict(updates)
    update_payload['updated_at'] = datetime.now().isoformat()

    if using_firestore():
        discovery_doc_ref(task_id).set(update_payload, merge=True)
        return

    allowed_columns = {
        'user_name', 'user_email', 'industry', 'form_data', 'status',
        'report_content', 'error_message', 'created_at', 'updated_at'
    }
    assignments = []
    values = []
    for key, value in update_payload.items():
        if key not in allowed_columns:
            continue
        assignments.append(f"{key} = ?")
        values.append(value)
    if not assignments:
        return
    values.append(str(task_id))
    conn = get_discovery_db()
    conn.execute(
        f"UPDATE discovery_tasks SET {', '.join(assignments)} WHERE task_id = ?",
        values
    )
    conn.commit()
    conn.close()

def get_discovery_task_record(task_id):
    if using_firestore():
        snapshot = discovery_doc_ref(task_id).get()
        if not snapshot.exists:
            return None
        data = snapshot.to_dict() or {}
        data.setdefault('task_id', str(task_id))
        return data

    conn = get_discovery_db()
    row = conn.execute("SELECT * FROM discovery_tasks WHERE task_id = ?", (task_id,)).fetchone()
    conn.close()
    return dict(row) if row else None

def list_recent_discovery_task_records(limit=30, status=None):
    limit = max(1, int(limit or 30))
    status = (status or '').strip().upper()

    if using_firestore():
        client = get_firestore_client()
        docs = client.collection(FIRESTORE_DISCOVERY_COLLECTION).order_by(
            'created_at',
            direction=firestore.Query.DESCENDING
        ).limit(max(limit * 5, 50)).stream()
        tasks = []
        for doc in docs:
            data = doc.to_dict() or {}
            if status and data.get('status') != status:
                continue
            data.setdefault('task_id', doc.id)
            tasks.append(data)
            if len(tasks) >= limit:
                break
        return tasks

    conn = get_discovery_db()
    try:
        if status:
            rows = conn.execute(
                "SELECT task_id, status, user_email, industry, created_at, updated_at, error_message "
                "FROM discovery_tasks WHERE status = ? ORDER BY created_at DESC LIMIT ?",
                (status, limit)
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT task_id, status, user_email, industry, created_at, updated_at, error_message "
                "FROM discovery_tasks ORDER BY created_at DESC LIMIT ?",
                (limit,)
            ).fetchall()
    finally:
        conn.close()
    return [dict(r) for r in rows]

def summarize_discovery_tasks(since_hours=24):
    since_hours = max(1, min(int(since_hours or 24), 24 * 30))
    since_ts = (datetime.now() - timedelta(hours=since_hours)).isoformat()

    if using_firestore():
        client = get_firestore_client()
        docs = client.collection(FIRESTORE_DISCOVERY_COLLECTION).order_by(
            'created_at',
            direction=firestore.Query.DESCENDING
        ).limit(1000).stream()
        by_status = {}
        active_count = 0
        latest = None
        for doc in docs:
            data = doc.to_dict() or {}
            created_at = data.get('created_at') or ''
            if created_at and created_at < since_ts:
                continue
            status = data.get('status') or 'UNKNOWN'
            by_status[status] = by_status.get(status, 0) + 1
            if status in ('PENDING', 'ANALYZING'):
                active_count += 1
            if latest is None:
                data.setdefault('task_id', doc.id)
                latest = {
                    'task_id': data.get('task_id'),
                    'status': data.get('status'),
                    'user_email': data.get('user_email'),
                    'created_at': data.get('created_at'),
                    'updated_at': data.get('updated_at'),
                }
        return {
            "since_hours": since_hours,
            "since_ts": since_ts,
            "active_count": active_count,
            "by_status": by_status,
            "latest": latest
        }

    conn = get_discovery_db()
    try:
        rows = conn.execute(
            "SELECT status, COUNT(*) AS count FROM discovery_tasks WHERE created_at >= ? GROUP BY status",
            (since_ts,)
        ).fetchall()
        active = conn.execute(
            "SELECT COUNT(*) AS count FROM discovery_tasks WHERE status IN ('PENDING','ANALYZING')"
        ).fetchone()
        latest = conn.execute(
            "SELECT task_id, status, user_email, created_at, updated_at FROM discovery_tasks ORDER BY created_at DESC LIMIT 1"
        ).fetchone()
    finally:
        conn.close()

    return {
        "since_hours": since_hours,
        "since_ts": since_ts,
        "active_count": int(active['count']) if active else 0,
        "by_status": {row['status']: row['count'] for row in rows},
        "latest": dict(latest) if latest else None
    }

def sync_discovery_task_backfill():
    if not using_firestore():
        return
    try:
        conn = get_discovery_db()
        rows = conn.execute(
            "SELECT task_id, user_name, user_email, industry, form_data, status, report_content, error_message, created_at, updated_at "
            "FROM discovery_tasks ORDER BY created_at DESC LIMIT 500"
        ).fetchall()
        conn.close()
        for row in rows:
            payload = dict(row)
            payload['user_email'] = normalize_email(payload.get('user_email'))
            discovery_doc_ref(payload['task_id']).set(payload, merge=True)
    except Exception as e:
        logger.error(f"Discovery task backfill error: {e}")

def sync_user_quota_backfill():
    try:
        now = datetime.now().isoformat()
        conn = get_db_connection()
        conn.execute(
            '''
            CREATE TABLE IF NOT EXISTS user_quota (
                email TEXT PRIMARY KEY,
                competitor_usage INTEGER NOT NULL DEFAULT 0,
                discovery_usage INTEGER NOT NULL DEFAULT 0,
                updated_at TEXT
            )
            '''
        )
        order_rows = conn.execute(
            "SELECT email, MAX(quota_usage) AS usage FROM orders WHERE email IS NOT NULL AND TRIM(email) != '' GROUP BY email"
        ).fetchall()
        quota_rows = conn.execute(
            "SELECT email, competitor_usage, discovery_usage FROM user_quota WHERE email IS NOT NULL AND TRIM(email) != ''"
        ).fetchall()
        for row in order_rows:
            email = normalize_email(row['email'])
            usage = int(row['usage'] or 0)
            existing = conn.execute("SELECT competitor_usage FROM user_quota WHERE email = ?", (email,)).fetchone()
            current = int(existing['competitor_usage']) if existing else 0
            conn.execute(
                '''
                INSERT INTO user_quota (email, competitor_usage, discovery_usage, updated_at)
                VALUES (?, ?, 0, ?)
                ON CONFLICT(email) DO UPDATE SET
                    competitor_usage = excluded.competitor_usage,
                    updated_at = excluded.updated_at
                ''',
                (email, max(current, usage), now)
            )
        conn.commit()
        conn.close()

        dconn = get_discovery_db()
        discovery_rows = dconn.execute(
            "SELECT user_email, COUNT(*) AS usage FROM discovery_tasks WHERE status IN ('COMPLETED','COMPLETED_NO_EMAIL') GROUP BY user_email"
        ).fetchall()
        dconn.close()

        conn = get_db_connection()
        merged_rows = {}
        for row in quota_rows:
            email = normalize_email(row['email'])
            if not email:
                continue
            merged_rows[email] = {
                'competitor_usage': int(row['competitor_usage'] or 0),
                'discovery_usage': int(row['discovery_usage'] or 0),
            }
        for row in discovery_rows:
            email = normalize_email(row['user_email'])
            usage = int(row['usage'] or 0)
            existing = conn.execute("SELECT discovery_usage FROM user_quota WHERE email = ?", (email,)).fetchone()
            current = int(existing['discovery_usage']) if existing else 0
            merged = max(current, usage)
            conn.execute(
                '''
                INSERT INTO user_quota (email, competitor_usage, discovery_usage, updated_at)
                VALUES (?, 0, ?, ?)
                ON CONFLICT(email) DO UPDATE SET
                    discovery_usage = excluded.discovery_usage,
                    updated_at = excluded.updated_at
                ''',
                (email, merged, now)
            )
            current_merged = merged_rows.get(email, {'competitor_usage': 0, 'discovery_usage': 0})
            current_merged['discovery_usage'] = max(current_merged['discovery_usage'], merged)
            merged_rows[email] = current_merged
        conn.commit()

        if using_firestore():
            for row in order_rows:
                email = normalize_email(row['email'])
                if not email:
                    continue
                current_merged = merged_rows.get(email, {'competitor_usage': 0, 'discovery_usage': 0})
                current_merged['competitor_usage'] = max(current_merged['competitor_usage'], int(row['usage'] or 0))
                merged_rows[email] = current_merged

            for email, usage_row in merged_rows.items():
                set_quota_usage(
                    email,
                    competitor_usage=usage_row['competitor_usage'],
                    discovery_usage=usage_row['discovery_usage'],
                )
        conn.commit()
        conn.close()
    except Exception as e:
        logger.error(f"User quota backfill error: {e}")

def set_quota_usage(email, competitor_usage=None, discovery_usage=None):
    normalized_email = normalize_email(email)
    if not normalized_email:
        return

    now = datetime.now().isoformat()
    competitor_usage = None if competitor_usage is None else max(0, int(competitor_usage))
    discovery_usage = None if discovery_usage is None else max(0, int(discovery_usage))

    if using_firestore():
        doc_ref = quota_doc_ref(normalized_email)
        snapshot = doc_ref.get()
        current = snapshot.to_dict() or {}
        payload = {
            'email': normalized_email,
            'updated_at': now,
            'competitor_usage': int(current.get('competitor_usage', 0)),
            'discovery_usage': int(current.get('discovery_usage', 0)),
        }
        if competitor_usage is not None:
            payload['competitor_usage'] = max(payload['competitor_usage'], competitor_usage)
        if discovery_usage is not None:
            payload['discovery_usage'] = max(payload['discovery_usage'], discovery_usage)
        doc_ref.set(payload, merge=True)
        return payload

    conn = get_db_connection()
    row = conn.execute(
        "SELECT competitor_usage, discovery_usage FROM user_quota WHERE email = ?",
        (normalized_email,)
    ).fetchone()
    current_competitor = int(row['competitor_usage']) if row else 0
    current_discovery = int(row['discovery_usage']) if row else 0
    conn.execute(
        '''
        INSERT INTO user_quota (email, competitor_usage, discovery_usage, updated_at)
        VALUES (?, ?, ?, ?)
        ON CONFLICT(email) DO UPDATE SET
            competitor_usage = excluded.competitor_usage,
            discovery_usage = excluded.discovery_usage,
            updated_at = excluded.updated_at
        ''',
        (
            normalized_email,
            max(current_competitor, competitor_usage or 0),
            max(current_discovery, discovery_usage or 0),
            now,
        )
    )
    conn.commit()
    conn.close()

def get_quota_status(email, feature):
    normalized_email = normalize_email(email)
    normalized_feature = normalize_quota_feature(feature)
    usage_key = 'discovery_usage' if normalized_feature == 'discovery' else 'competitor_usage'

    if not normalized_email:
        return {
            "email": "",
            "feature": normalized_feature,
            "usage": 0,
            "remaining": FREE_QUOTA_LIMIT,
            "allowed": False
        }

    if using_firestore():
        snapshot = quota_doc_ref(normalized_email).get()
        row = snapshot.to_dict() if snapshot.exists else {}
        usage = int((row or {}).get(usage_key, 0) or 0)
    else:
        conn = get_db_connection()
        row = conn.execute(
            "SELECT competitor_usage, discovery_usage FROM user_quota WHERE email = ?",
            (normalized_email,)
        ).fetchone()
        conn.close()
        usage = int(row[usage_key]) if row else 0

    remaining = max(0, FREE_QUOTA_LIMIT - usage)
    return {
        "email": normalized_email,
        "feature": normalized_feature,
        "usage": usage,
        "remaining": remaining,
        "allowed": usage < FREE_QUOTA_LIMIT
    }

def increment_quota_usage(email, feature, amount=1):
    normalized_email = normalize_email(email)
    normalized_feature = normalize_quota_feature(feature)
    if not normalized_email:
        return None

    now = datetime.now().isoformat()
    amount = int(amount or 0)
    if using_firestore():
        doc_ref = quota_doc_ref(normalized_email)

        @firestore.transactional
        def update_in_transaction(transaction):
            snapshot = doc_ref.get(transaction=transaction)
            current = snapshot.to_dict() or {}
            competitor_usage = int(current.get('competitor_usage', 0) or 0)
            discovery_usage = int(current.get('discovery_usage', 0) or 0)
            if normalized_feature == 'discovery':
                discovery_usage += amount
            else:
                competitor_usage += amount
            payload = {
                'email': normalized_email,
                'competitor_usage': max(0, competitor_usage),
                'discovery_usage': max(0, discovery_usage),
                'updated_at': now,
            }
            transaction.set(doc_ref, payload, merge=True)
            return payload

        return update_in_transaction(get_firestore_client().transaction())

    usage_column = 'discovery_usage' if normalized_feature == 'discovery' else 'competitor_usage'
    conn = get_db_connection()
    row = conn.execute(
        "SELECT competitor_usage, discovery_usage FROM user_quota WHERE email = ?",
        (normalized_email,)
    ).fetchone()
    competitor_usage = int(row['competitor_usage']) if row else 0
    discovery_usage = int(row['discovery_usage']) if row else 0
    if usage_column == 'competitor_usage':
        competitor_usage += amount
    else:
        discovery_usage += amount
    payload = {
        'email': normalized_email,
        'competitor_usage': max(0, competitor_usage),
        'discovery_usage': max(0, discovery_usage),
        'updated_at': now,
    }
    conn.execute(
        '''
        INSERT INTO user_quota (email, competitor_usage, discovery_usage, updated_at)
        VALUES (?, ?, ?, ?)
        ON CONFLICT(email) DO UPDATE SET
            competitor_usage = excluded.competitor_usage,
            discovery_usage = excluded.discovery_usage,
            updated_at = excluded.updated_at
        ''',
        (
            normalized_email,
            payload['competitor_usage'],
            payload['discovery_usage'],
            payload['updated_at'],
        )
    )
    conn.commit()
    conn.close()
    return payload

sync_user_quota_backfill()
sync_discovery_task_backfill()
logger.info(f"Persistence backend: {'firestore' if using_firestore() else 'sqlite'}")

def _generate_discovery_report(form_data):
    """Generate AI report for product discovery."""
    if not OPENAI_API_KEY:
        return "ERROR: API Key not configured. Add OPENAI_API_KEY or OPENROUTER_API_KEY in the server environment variables."

    model_to_use = form_data.get('ai_model', OPENAI_MODEL)
    report_lang = form_data.get('report_language', 'zh')
    lang_instruction = """
    [语言要求]:
    1. 报告主体必须使用简体中文撰写。
    2. 亚马逊关键词、产品术语、用户原话引用保留原始语言（英文/德文/日文等）。
    """ if report_lang == 'zh' else """
    [Language Requirement]:
    1. The report must be written entirely in English.
    2. Keep Amazon keywords and product terms in their original form.
    """

    prompt = f"""你是一位顶级的亚马逊选品与竞品分析专家，曾为数百个卖家撰写过单份售价数千元的深度报告。请基于以下市场发现请求，生成一份**专业级全品类深度分析报告**，结构与深度对标 FlowAI Agent 付费竞品分析报告标准。

【字数硬性要求】报告正文必须达到 8000 汉字以上（英文报告 4000+ words），务必写足写满，不要提前结束。每个表格需有 4-8 行实质内容，每个小节需展开论述（每段 80-150 字），切忌泛泛而谈。

[输入参数]
- 主类目: {form_data.get('category_main')}
- 类目路径: {form_data.get('category_path')}
- 产品关键词: {form_data.get('keywords')}
- 目标站点: {form_data.get('marketplace')}
- 分析重点: {', '.join(form_data.get('focus_areas', []))}
- 目标价格区间: ${form_data.get('target_price_min') or '—'} - ${form_data.get('target_price_max') or '—'}
- 参考ASIN (如有): {', '.join(form_data.get('reference_asins', [])) or '无'}
- 个性化要求: {form_data.get('custom_prompt') or '无'}
{lang_instruction}

[报告结构 - 严格按以下六大部分撰写，每部分需有实质数据与可执行建议]

---

# 全品类深度分析报告: [品类名称]

## 一、市场吸引力综合评估
### 1.1 市场评分卡
用表格呈现，每维度 1-10 分，并附 80-120 字详细分析：
| 维度 | 评分(1-10) | 详细分析 |
| 需求强度 | X | [分析该品类的需求稳定性、复购性、季节性、用户群体规模] |
| 竞争壁垒 | X | [分析品牌垄断度、技术门槛、供应链难度、新进入者空间] |
| 利润潜力 | X | [分析成本结构、毛利率空间、LTV、价格战风险] |
| 进入难度 | X | [分析资金需求、供应链、认证、初期获客难度] |

**综合市场吸引力评分: X.X/10**
**综合判断: ENTER / CAUTIOUS ENTER / NO-GO**
**详细理由:** [200-300字，结合上述维度给出进入或谨慎进入的具体理由与风险提示]

### 1.2 价格生态系统深度分析
- **价格影响因素矩阵** (表格): 因素 | 对价格的影响 | 说明
- **推演的价格带分布**: 引流款区间、主流款区间、高端款区间，每区间需说明特征、竞争程度、策略建议
- **多包装规格价格策略建议** (表格): 规格 | 建议定价 | 单片/单位价格 | 目标用户 | 战略目的
- **建议切入价位**: 给出具体 $XX.XX，并详述 5-6 条定价理由（心理定价、成本结构、竞争定位、促销空间等）

---

## 二、用户群体深度画像
### 2.1 核心用户画像表
| 画像编号 | 用户类型 | 年龄 | 职业特征 | 购买动机 | 价格敏感度 | 占比估计 |
[4-5 行，每类用户有具体描述]

### 2.2 场景化需求分析表
| 优先级 | 使用场景 | 详细描述 | 转化驱动因素 | 评论证据/推理 |
[5-8 行，每个场景 50-80 字描述，转化驱动因素 2-3 条]

### 2.3 用户痛点深度挖掘
| 排名 | 痛点 | 占比 | 根因分析 | 影响用户群 | 解决难度 | 差异化价值 |
[5-8 个痛点，每个有根因分析与对策方向]

### 2.4 用户决策路径分析
- **专业用户/批量采购者** 的购买决策流程 (5-8 步)
- **个人消费者/DIY 用户** 的购买决策流程 (5-8 步)
- **关键决策节点** 与 Listing 优化对应建议

---

## 三、产品与技术趋势
### 3.1 材质/设计演进表
| 阶段 | 主流材质/技术 | 核心功能 | 优势 | 劣势 | 市场状态 |
[过去、当前、未来趋势，3-5 行]

### 3.2 规格/型号趋势分析
针对该品类的关键规格维度（如粒度、尺寸、功率等），分析各规格的市场需求趋势

### 3.3 爆款基因解码
- **必备功能 (Table Stakes)**: 表格列出 5-6 项，不具备就出局
- **加分功能 (Points of Difference)**: 表格列出 4-5 项，具备可脱颖而出
- **设计语言**: 色彩、信息层级、主图要求、描述风格
- **定价区间**: 基于成功产品的价格带分析

### 3.4 技术创新机会
| 创新方向 | 技术描述 | 解决的问题 | 可行性 | 潜在影响 |
[3-5 个创新方向]

---

## 四、竞争格局详解
### 4.1 品牌垄断度分析
- 市场结构评估 (头部集中度、中小品牌空间、白牌比例、价格竞争度)
- 头部品牌识别 (3-5 个典型品牌及定位)
- 市场格局判断: 垄断型/寡头型/分散竞争型
- 新卖家切入可能性: 高/中/低，及理由

### 4.2 竞品深度解剖
选取 1-2 个典型竞品（可基于品类知识虚构或使用"假设竞品A/B"），每个包含：
- 基本信息 (价格、评分、评论数、核心规格)
- 核心优势 (从品类共性推断，3-5 条，每条附"评论证据"风格的说明)
- 核心劣势 (3-4 条，含根因分析)
- 市场定位与竞争策略建议

### 4.3 竞品对比矩阵与市场空白分析
- 维度对比表格 (价格、评分、材料、功能、包装、目标用户等)
- **识别的市场空白和机会** (表格): 机会编号 | 市场空白描述 | 目标用户 | 产品方案 | 预期定价 | 差异化价值
[4-6 个具体机会]
- **推荐的市场进入策略**: 三阶段策略 (0-3月、3-6月、6-12月)

---

## 五、蓝海机会与差异化策略
### 5.1 未被满足的需求深度挖掘
| 排名 | 未满足需求 | 严重程度 | 证据/推理 | 影响用户群 | 解决难度 | 差异化价值 | 实施优先级 |
[5-8 项，每项有实质分析]

### 5.2 差异化策略矩阵
| 策略类型 | 具体方案 | 实施成本 | 实施周期 | 预期效果 | 风险 | 优先级 |
[6-10 项，涵盖产品创新、包装创新、内容创新、服务创新、定位创新、SKU 策略]

### 5.3 推荐的差异化组合策略
- **阶段一 (0-2 月)**: 低成本快速差异化，3-4 项具体措施
- **阶段二 (2-4 月)**: 产品线扩展，2-3 项
- **阶段三 (4-6 月)**: 技术差异化，2-3 项

### 5.4 差异化价值主张设计
- 核心价值主张 (一句话)
- 支撑点 (4-5 条)
- 针对不同用户群的定制化价值主张 (表格)

---

## 六、进入策略与执行计划
### 6.1 理想产品规格定义
- 维度 | 规格要求 | 理由 | 供应商要求 (表格)
- 目标成本结构 (采购、运输、FBA、佣金、广告、目标净利润)
- 目标售价与定价理由

### 6.2 供应链与质量控制计划
- 供应商筛选标准 (6-8 项)
- 供应商开发流程 (4-5 阶段)
- 首批订单建议 (数量、分配、总投资)
- 质量控制流程 (生产前/中/后/到货/上架)

### 6.3 风险雷达与缓解措施
| 风险类型 | 具体风险 | 概率 | 影响 | 缓解措施 | 应急预案 |
[8-12 项，涵盖市场、供应链、运营、合规、财务、平台风险]

### 6.4 30-60-90 天执行计划
- **第 1-30 天 (筹备期)**: 按周分解，每周 3-5 个具体任务，含关键产出
- **第 31-60 天 (生产和启动期)**: 同上
- **第 61-90 天 (验证和优化期)**: 同上
- 每阶段需有明确的成功标准与决策点 (验证成功/需优化/考虑退出)

---

**重要**: 整份报告必须不少于 8000 汉字。基于品类知识进行合理推断，数据可标注为"基于行业经验推断"或"合理估算"。每个表格和列表都要有实质信息，避免空洞概括。"""

    def call_model(messages):
        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://flowaiagent.com",
            "X-Title": "FlowAI Agent"
        }
        payload = {
            "model": model_to_use,
            "messages": messages,
            "temperature": 0.6,
            "max_tokens": 16384
        }
        logger.info(f"Discovery: Generating report using {model_to_use} (target 8000+ chars)")
        response = requests.post(f"{OPENAI_BASE_URL}/chat/completions", json=payload, headers=headers, timeout=300)
        if not response.ok:
            err_body = response.text[:500] if response.text else "(empty)"
            logger.error(f"Discovery OpenRouter API Error: status={response.status_code}, body={err_body}")

            if response.status_code == 401:
                return (
                    "ERROR: OpenRouter rejected the API key (HTTP 401). "
                    "Please verify OPENAI_API_KEY or OPENROUTER_API_KEY in the server environment. "
                    f"Details: {err_body}",
                    None
                )
            if response.status_code == 402:
                return (
                    "ERROR: OpenRouter credits are insufficient (HTTP 402). "
                    "Please recharge openrouter.ai/settings/credits. "
                    f"Details: {err_body}",
                    None
                )
            if response.status_code == 429:
                return (
                    "ERROR: OpenRouter rate limit reached (HTTP 429). "
                    "Please retry shortly or switch to another model. "
                    f"Details: {err_body}",
                    None
                )

            return (
                f"ERROR: OpenRouter API failed (HTTP {response.status_code}). "
                f"Details: {err_body}",
                None
            )
        data = response.json()
        choice = data['choices'][0]
        content = choice['message']['content']
        finish_reason = choice.get('finish_reason')
        return content, finish_reason

    def report_looks_complete(content):
        normalized = re.sub(r'\s+', '', content)
        completion_markers = [
            '30-60-90天执行计划',
            '第61-90天',
            '6.4',
            '##六、进入策略与执行计划',
            '###6.4',
            '第31-60天',
        ]
        return any(marker in normalized for marker in completion_markers)

    try:
        messages = [{"role": "user", "content": prompt}]
        content, finish_reason = call_model(messages)
        if isinstance(content, str) and content.startswith("ERROR:"):
            return content

        combined = content
        logger.info(f"Discovery: Report generated chunk 1, {len(content)} chars, finish_reason={finish_reason}")

        for attempt in range(2):
            if finish_reason != 'length' and report_looks_complete(combined):
                break
            continuation_prompt = (
                "Continue the same report from exactly where you stopped. "
                "Do not repeat previous sections. Resume from the next unfinished heading and finish the remaining sections."
            )
            messages = [
                {"role": "user", "content": prompt},
                {"role": "assistant", "content": combined},
                {"role": "user", "content": continuation_prompt}
            ]
            next_chunk, finish_reason = call_model(messages)
            if isinstance(next_chunk, str) and next_chunk.startswith("ERROR:"):
                break
            if not next_chunk.strip():
                break
            combined += "\n\n" + next_chunk.strip()
            logger.info(
                f"Discovery: Report continuation chunk {attempt + 2}, {len(next_chunk)} chars, "
                f"total={len(combined)}, finish_reason={finish_reason}"
            )

        logger.info(f"Discovery: Report generated, {len(combined)} chars total")
        return combined
    except requests.exceptions.RequestException as e:
        err_detail = str(e)
        if hasattr(e, 'response') and e.response is not None:
            err_detail = f"HTTP {e.response.status_code}: {e.response.text[:300] if e.response.text else 'no body'}"
        logger.error(f"Discovery AI Request Error: {err_detail}")
        return f"ERROR: API request failed. {err_detail}"
    except Exception as e:
        logger.error(f"Discovery AI Error: {e}")
        return f"ERROR: Failed to generate report. {str(e)}"

def _record_discovery_usage(email):
    """Increment quota_usage for discovery (选品分析) completion."""
    if not email:
        return
    try:
        increment_quota_usage(email, 'discovery', 1)
        logger.info(f"Discovery: Recorded discovery usage for {normalize_email(email)}")
    except Exception as e:
        logger.error(f"Discovery record_usage error: {e}")

def _render_inline_markdown(text):
    escaped = escape(text)
    escaped = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', escaped)
    escaped = re.sub(r'`([^`]+)`', r'<code>\1</code>', escaped)
    return escaped

def _markdown_to_basic_html(markdown_text):
    parts = []
    list_open = False
    table_rows = []

    def flush_list():
        nonlocal list_open
        if list_open:
            parts.append('</ul>')
            list_open = False

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return

        def is_separator(row):
            return all(cell and set(cell) <= set('-: ') for cell in row)

        rows = table_rows[:]
        table_rows = []
        if len(rows) >= 2 and is_separator(rows[1]):
            header = rows[0]
            body = rows[2:]
        else:
            header = rows[0]
            body = rows[1:]

        parts.append('<table style="width:100%;border-collapse:collapse;margin:16px 0;">')
        parts.append('<thead><tr>')
        for cell in header:
            parts.append(
                f'<th style="border:1px solid #dbe3f0;padding:8px 10px;background:#f8fafc;text-align:left;">{_render_inline_markdown(cell)}</th>'
            )
        parts.append('</tr></thead>')
        if body:
            parts.append('<tbody>')
            for row in body:
                parts.append('<tr>')
                for cell in row:
                    parts.append(
                        f'<td style="border:1px solid #dbe3f0;padding:8px 10px;vertical-align:top;">{_render_inline_markdown(cell)}</td>'
                    )
                parts.append('</tr>')
            parts.append('</tbody>')
        parts.append('</table>')

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith('|') and stripped.endswith('|'):
            flush_list()
            cells = [cell.strip() for cell in stripped.strip('|').split('|')]
            table_rows.append(cells)
            continue
        flush_table()

        if not stripped:
            flush_list()
            continue

        if stripped.startswith('# '):
            flush_list()
            parts.append(f'<h1 style="font-size:24px;margin:28px 0 14px;">{_render_inline_markdown(stripped[2:])}</h1>')
        elif stripped.startswith('## '):
            flush_list()
            parts.append(f'<h2 style="font-size:20px;margin:24px 0 12px;">{_render_inline_markdown(stripped[3:])}</h2>')
        elif stripped.startswith('### '):
            flush_list()
            parts.append(f'<h3 style="font-size:17px;margin:20px 0 10px;">{_render_inline_markdown(stripped[4:])}</h3>')
        elif stripped.startswith('- ') or stripped.startswith('* '):
            if not list_open:
                parts.append('<ul style="margin:12px 0 12px 22px;padding:0;">')
                list_open = True
            parts.append(f'<li style="margin:6px 0;">{_render_inline_markdown(stripped[2:])}</li>')
        else:
            flush_list()
            parts.append(f'<p style="margin:12px 0;line-height:1.75;">{_render_inline_markdown(stripped)}</p>')

    flush_list()
    flush_table()
    return ''.join(parts)

def _build_discovery_email_preview(report_content, max_lines=18, max_chars=2200):
    lines = []
    total = 0
    for raw_line in report_content.splitlines():
        stripped = raw_line.strip()
        if not stripped or stripped.startswith('|'):
            continue
        if set(stripped) <= set('-:| '):
            continue
        cleaned = re.sub(r'^[#*\-\d\.\s]+', '', stripped).strip()
        if not cleaned:
            continue
        lines.append(cleaned)
        total += len(cleaned)
        if len(lines) >= max_lines or total >= max_chars:
            break
    return '\n'.join(lines)

def _strip_markdown_markers(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text.strip()

def _report_to_docx_bytes(report_content, title_text):
    document = Document()
    normal_style = document.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    normal_style.font.size = Pt(10.5)

    table_rows = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return

        def is_separator(row):
            return all(cell and set(cell) <= set('-: ') for cell in row)

        rows = table_rows[:]
        table_rows = []
        if len(rows) >= 2 and is_separator(rows[1]):
            header = rows[0]
            body = rows[2:]
        else:
            header = rows[0]
            body = rows[1:]

        table = document.add_table(rows=1 + len(body), cols=len(header))
        table.style = 'Table Grid'
        for idx, cell in enumerate(header):
            table.rows[0].cells[idx].text = _strip_markdown_markers(cell)
        for row_idx, row in enumerate(body, start=1):
            for col_idx, cell in enumerate(row):
                table.rows[row_idx].cells[col_idx].text = _strip_markdown_markers(cell)
        document.add_paragraph('')

    if title_text:
        document.add_heading(_strip_markdown_markers(title_text), level=0)

    for raw_line in report_content.splitlines():
        stripped = raw_line.strip()
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [cell.strip() for cell in stripped.strip('|').split('|')]
            table_rows.append(cells)
            continue
        flush_table()
        if not stripped:
            continue
        if stripped.startswith('# '):
            document.add_heading(_strip_markdown_markers(stripped[2:]), level=1)
        elif stripped.startswith('## '):
            document.add_heading(_strip_markdown_markers(stripped[3:]), level=2)
        elif stripped.startswith('### '):
            document.add_heading(_strip_markdown_markers(stripped[4:]), level=3)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            document.add_paragraph(_strip_markdown_markers(stripped[2:]), style='List Bullet')
        else:
            document.add_paragraph(_strip_markdown_markers(stripped))

    flush_table()

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()

def _send_discovery_email(to_email, user_name, report_content):
    """Send discovery report to user email."""
    if not SMTP_USER or not SMTP_PASSWORD:
        logger.warning("Discovery: SMTP not configured. Skipping email.")
        return False
    msg = MIMEMultipart('mixed')
    msg['From'] = SENDER_EMAIL
    msg['To'] = to_email
    msg['Subject'] = f"Your Product Discovery Report for {user_name}"

    safe_name = re.sub(r'[^A-Za-z0-9_-]+', '_', (user_name or 'user')).strip('_') or 'user'
    preview_text = _build_discovery_email_preview(report_content)
    report_docx = _report_to_docx_bytes(report_content, "Product Discovery Report")

    plain_body = (
        f"Hello {user_name},\n\n"
        "Your Amazon Product Discovery report is ready.\n\n"
        "For better readability and to avoid email clipping, the full report is attached as a DOCX document.\n\n"
        "Preview:\n"
        "----------------------------------------\n"
        f"{preview_text}\n\n"
        "Attachments:\n"
        f"- product_discovery_report_{safe_name}.docx\n"
    )

    html_body = f"""
    <html>
      <body style="margin:0;padding:24px;background:#f3f6fb;color:#0f172a;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;">
        <div style="max-width:840px;margin:0 auto;background:#ffffff;border:1px solid #e2e8f0;border-radius:16px;padding:32px;">
          <h1 style="margin:0 0 12px;font-size:28px;">Product Discovery Report</h1>
          <p style="margin:0 0 12px;line-height:1.7;">Hi {escape(user_name)}, your report is ready.</p>
          <p style="margin:0 0 18px;line-height:1.7;">
            To avoid long-email clipping, the <strong>full report is attached</strong> as a DOCX document that opens cleanly in Word / Google Docs.
            The email body only shows a readable preview.
          </p>
          <div style="background:#f8fafc;border:1px solid #dbe3f0;border-radius:12px;padding:18px 20px;margin:20px 0;">
            {_markdown_to_basic_html(preview_text)}
          </div>
          <p style="margin:18px 0 0;color:#475569;">Attachments:</p>
          <ul style="margin:8px 0 0 20px;line-height:1.8;color:#475569;">
            <li>product_discovery_report_{safe_name}.docx</li>
          </ul>
        </div>
      </body>
    </html>
    """

    alternative = MIMEMultipart('alternative')
    alternative.attach(MIMEText(plain_body, 'plain', 'utf-8'))
    alternative.attach(MIMEText(html_body, 'html', 'utf-8'))
    msg.attach(alternative)

    docx_attachment = MIMEApplication(
        report_docx,
        _subtype='vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    docx_attachment.add_header('Content-Disposition', 'attachment', filename=f'product_discovery_report_{safe_name}.docx')
    msg.attach(docx_attachment)

    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SMTP_USER, SMTP_PASSWORD)
        server.send_message(msg)
        server.quit()
        logger.info(f"Discovery: Email sent to {to_email}")
        return True
    except Exception as e:
        logger.error(f"Discovery Email Error: {e}")
        return False

def _discovery_worker(task_id, user_name, user_email, form_data):
    """Background worker for discovery report generation."""
    try:
        update_discovery_task_record(task_id, status='ANALYZING')

        report = _generate_discovery_report(form_data)
        if report.startswith("ERROR:"):
            raise Exception(report)

        email_sent = _send_discovery_email(user_email, user_name, report)

        update_discovery_task_record(
            task_id,
            status='COMPLETED' if email_sent else 'COMPLETED_NO_EMAIL',
            report_content=report,
            error_message=''
        )
        _record_discovery_usage(user_email)
        logger.info(f"Discovery task {task_id} completed.")
    except Exception as e:
        logger.error(f"Discovery Worker Error for {task_id}: {e}")
        update_discovery_task_record(task_id, status='FAILED', error_message=str(e))

@app.route('/api/discovery/submit', methods=['POST'])
def discovery_submit():
    """Product Discovery: submit analysis task."""
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400

    task_id = str(uuid.uuid4())
    user_name = data.get('user_name', 'User')
    user_email = normalize_email(data.get('user_email', ''))
    industry = data.get('industry', '')

    if not user_email:
        return jsonify({"error": "Email is required"}), 400

    quota = get_quota_status(user_email, 'discovery')
    if not quota['allowed']:
        return jsonify({
            "error": "Discovery free quota exceeded",
            **quota
        }), 403

    create_discovery_task_record(task_id, user_name, user_email, industry, data)

    thread = threading.Thread(target=_discovery_worker, args=(task_id, user_name, user_email, data))
    thread.start()

    logger.info(f"Discovery task {task_id} submitted for {user_email}")
    return jsonify({"task_id": task_id, "status": "PENDING"}), 202

@app.route('/api/discovery/status', methods=['GET'])
def discovery_status():
    """Product Discovery: poll task status."""
    task_id = request.args.get('task_id')
    if not task_id:
        return jsonify({"error": "task_id required"}), 400

    row = get_discovery_task_record(task_id)
    if not row:
        return jsonify({"error": "Task not found"}), 404

    return jsonify({
        "status": row.get('status'),
        "error_message": (
            "当前分析任务未能完成，请稍后重试或联系客服。"
            if row.get('status') == 'FAILED' and row.get('error_message')
            else row.get('error_message')
        ),
        "updated_at": row.get('updated_at')
    })

@app.route('/api/discovery/admin/summary', methods=['GET'])
def discovery_admin_summary():
    """Admin: summarize recent discovery task statuses."""
    if not is_admin_request():
        return jsonify({"error": "Unauthorized"}), 403

    try:
        since_hours = int(request.args.get('since_hours', 24))
    except (TypeError, ValueError):
        since_hours = 24
    since_hours = max(1, min(since_hours, 24 * 30))

    return jsonify(summarize_discovery_tasks(since_hours))


@app.route('/api/discovery/admin/recent', methods=['GET'])
def discovery_admin_recent():
    """Admin: list recent discovery tasks (optionally filter by status)."""
    if not is_admin_request():
        return jsonify({"error": "Unauthorized"}), 403

    try:
        limit = int(request.args.get('limit', 30))
    except (TypeError, ValueError):
        limit = 30
    limit = max(1, min(limit, 200))

    status = (request.args.get('status') or '').strip().upper()
    allowed_status = {'PENDING', 'ANALYZING', 'COMPLETED', 'COMPLETED_NO_EMAIL', 'FAILED'}
    if status and status not in allowed_status:
        return jsonify({"error": "Invalid status filter"}), 400

    return jsonify({
        "limit": limit,
        "status": status or None,
        "tasks": list_recent_discovery_task_records(limit=limit, status=status or None)
    })

@app.route('/')
def home():
    """Serve the Main Website Homepage"""
    return send_from_directory('.', 'index.html')

@app.route('/index.html')
def redirect_index_to_root():
    """SEO: unify homepage URL to /"""
    return redirect('/', code=301)

@app.route('/create.html')
def redirect_create_to_create_analysis():
    """SEO: old create URL moved to create-analysis"""
    return redirect('/create-analysis.html', code=301)

@app.route('/api/check_quota', methods=['POST'])
def check_quota():
    """Check user quota for competitor/discovery (2 free uses each)."""
    data = request.json or {}
    email = normalize_email(data.get('email'))
    feature = normalize_quota_feature(data.get('feature'))
    
    if not email:
        return jsonify({"allowed": False, "error": "Email required"}), 400

    try:
        return jsonify(get_quota_status(email, feature))
    except Exception as e:
        logger.error(f"Quota Check Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/quota', methods=['GET'])
def quota_summary():
    """Lightweight quota summary for both features.

    Returns: { competitor_remaining: 1, discovery_remaining: 2, competitor_usage: 1, discovery_usage: 0, limit: 2 }
    """
    email = normalize_email(request.args.get('email'))
    if not email:
        return jsonify({"error": "email is required"}), 400
    try:
        competitor = get_quota_status(email, 'competitor')
        discovery = get_quota_status(email, 'discovery')
        return jsonify({
            "email": email,
            "limit": FREE_QUOTA_LIMIT,
            "competitor_usage": competitor.get("usage", 0),
            "discovery_usage": discovery.get("usage", 0),
            "competitor_remaining": competitor.get("remaining", 0),
            "discovery_remaining": discovery.get("remaining", 0),
        })
    except Exception as e:
        logger.error(f"Quota Summary Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/check_status', methods=['GET'])
def check_status():
    """Frontend Endpoint to Poll Status"""
    order_id = parse_order_id(request.args.get('order_id'))
    
    try:
        conn = get_db_connection()
        if order_id is None:
            conn.close()
            return jsonify({"error": "order_id is required"}), 400

        row = conn.execute('SELECT status, updated_at FROM orders WHERE id = ?', (order_id,)).fetchone()
             
        conn.close()
        
        if row:
            return jsonify(dict(row))
        else:
            return jsonify({"status": "NOT_FOUND"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/create_order', methods=['POST'])
def create_order():
    """Create a paid order that can later be manually confirmed by admin."""
    data = request.get_json(silent=True) or {}
    order_data = data.get('order_data', {})

    if not isinstance(order_data, dict):
        return jsonify({"error": "order_data must be an object"}), 400

    email = str(order_data.get('user_email') or data.get('email') or '').strip().lower()
    if not email:
        return jsonify({"error": "user email is required"}), 400

    updated_at = datetime.now().isoformat()
    conn = get_db_connection()
    created_order_id = None
    try:
        for _ in range(5):
            candidate_order_id = int(f"{int(time.time() * 1000)}{random.randint(100, 999)}")
            try:
                conn.execute(
                    'INSERT INTO orders (id, status, updated_at, order_data, email) VALUES (?, ?, ?, ?, ?)',
                    (candidate_order_id, 'PENDING', updated_at, json.dumps(order_data), email)
                )
                conn.commit()
                created_order_id = candidate_order_id
                break
            except sqlite3.IntegrityError:
                continue

        if created_order_id is None:
            return jsonify({"error": "Failed to allocate order_id"}), 500

        logger.info(f"Created order {created_order_id} for {email}")
        return jsonify({"order_id": str(created_order_id), "status": "PENDING"}), 201
    except Exception as e:
        logger.error(f"Create Order Error: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()

# n8n Webhook URL for report generation (override in env for production)
N8N_WEBHOOK_URL = os.environ.get(
    'N8N_WEBHOOK_URL',
    'https://tony4927.app.n8n.cloud/webhook/1573cd32-8e6a-46ac-9d74-1e6f7c9ea5e7',
)

def _normalize_n8n_payload(order_data):
    normalized = dict(order_data or {})
    user_name = normalized.get('user_name') or normalized.get('userName') or normalized.get('name') or ''
    user_email = normalize_email(
        normalized.get('user_email')
        or normalized.get('userEmail')
        or normalized.get('email')
        or ''
    )
    main_asins = normalize_amazon_asin_list(
        normalized.get('main_asins') or normalized.get('mainAsins') or []
    )
    competitor_asins = normalize_amazon_asin_list(
        normalized.get('competitor_asins') or normalized.get('competitorAsins') or []
    )
    custom_prompt = normalized.get('custom_prompt') or normalized.get('customPrompt') or ''
    site_count = normalized.get('reference_site_count') or normalized.get('referenceSiteCount') or 10
    youtube_count = normalized.get('reference_youtube_count') or normalized.get('referenceYoutubeCount') or 10

    normalized.update({
        'source': normalized.get('source') or 'payment-success',
        'analysis_type': normalized.get('analysis_type') or 'competitor_analysis',
        'user_name': user_name,
        'userName': user_name,
        'name': user_name,
        'user_email': user_email,
        'userEmail': user_email,
        'email': user_email,
        'main_asins': main_asins,
        'mainAsins': main_asins,
        'competitor_asins': competitor_asins,
        'competitorAsins': competitor_asins,
        'custom_prompt': custom_prompt,
        'customPrompt': custom_prompt,
        'reference_site_count': site_count,
        'referenceSiteCount': site_count,
        'reference_youtube_count': youtube_count,
        'referenceYoutubeCount': youtube_count,
    })
    return normalized

def _post_n8n_webhook(order_data):
    normalized = _normalize_n8n_payload(order_data)
    if not N8N_WEBHOOK_URL:
        logger.error("N8N_WEBHOOK_URL is not configured")
        raise ValueError("N8N_WEBHOOK_URL is not configured")
    logger.info(f"Triggering n8n webhook with data: {normalized}")
    response = requests.post(
        N8N_WEBHOOK_URL,
        json=normalized,
        headers={'Content-Type': 'application/json'},
        timeout=30
    )
    logger.info(f"n8n webhook response: {response.status_code} - {response.text[:200]}")
    return normalized, response

def trigger_n8n_webhook(order_data):
    """Trigger n8n webhook to generate report"""
    try:
        _, response = _post_n8n_webhook(order_data)
        return response.status_code == 200
    except Exception as e:
        logger.error(f"Error triggering n8n webhook: {e}")
        return False

@app.route('/api/competitor/submit', methods=['POST'])
def competitor_submit():
    """Server-side submission proxy for competitor analysis with quota enforcement."""
    data = request.get_json(silent=True) or {}
    email = normalize_email(
        data.get('user_email') or data.get('userEmail') or data.get('email')
    )
    if not email:
        return jsonify({"error": "Email required"}), 400

    quota = get_quota_status(email, 'competitor')
    if not quota['allowed']:
        return jsonify({
            "error": "Competitor free quota exceeded",
            **quota
        }), 403

    try:
        normalized, response = _post_n8n_webhook(data)
        if response.status_code != 200:
            return jsonify({
                "error": "Webhook failed",
                "status_code": response.status_code,
                "body": response.text[:300]
            }), 502

        increment_quota_usage(email, 'competitor', 1)
        quota_after = get_quota_status(email, 'competitor')
        return jsonify({
            "success": True,
            "feature": "competitor",
            "usage": quota_after['usage'],
            "remaining": quota_after['remaining'],
            "source": normalized.get('source') or 'create-analysis'
        })
    except Exception as e:
        logger.error(f"Competitor submit error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/record_usage', methods=['POST'])
def record_usage():
    """Increment user usage count for a specific feature (admin-only; was public and abusable)."""
    if not is_admin_request():
        return jsonify({"error": "Unauthorized"}), 403
    data = request.json or {}
    email = normalize_email(data.get('email'))
    feature = normalize_quota_feature(data.get('feature'))
    
    if not email:
        return jsonify({"error": "Email required"}), 400

    try:
        increment_quota_usage(email, feature, 1)
        quota = get_quota_status(email, feature)
        return jsonify({"success": True, **quota})
    except Exception as e:
        logger.error(f"Record Usage Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/quota', methods=['GET', 'POST'])
def admin_quota():
    """Admin: inspect or set quota usage for a specific email."""
    if not is_admin_request():
        return jsonify({"error": "Unauthorized"}), 403

    if request.method == 'GET':
        email = normalize_email(request.args.get('email'))
        feature = normalize_quota_feature(request.args.get('feature'))
        if not email:
            return jsonify({"error": "email is required"}), 400
        return jsonify(get_quota_status(email, feature))

    data = request.get_json(silent=True) or {}
    email = normalize_email(data.get('email'))
    feature = normalize_quota_feature(data.get('feature'))
    usage = data.get('usage')
    if not email:
        return jsonify({"error": "email is required"}), 400
    if usage is None:
        return jsonify({"error": "usage is required"}), 400

    try:
        usage = max(0, int(usage))
    except (TypeError, ValueError):
        return jsonify({"error": "usage must be an integer"}), 400
    if feature == 'discovery':
        set_quota_usage(email, discovery_usage=usage)
    else:
        set_quota_usage(email, competitor_usage=usage)
    return jsonify(get_quota_status(email, feature))

@app.route('/api/update_status', methods=['POST'])
def update_status():
    """Admin Endpoint to Update Status and trigger webhook on SUCCESS"""
    if not is_admin_request():
        return jsonify({"error": "Unauthorized"}), 403

    data = request.get_json(silent=True) or {}
    new_status = data.get('status')
    if new_status not in ['SUCCESS', 'FAILED', 'PENDING']:
        return jsonify({"error": "Invalid status"}), 400

    order_id = parse_order_id(data.get('order_id'))
    if order_id is None:
        return jsonify({"error": "order_id is required"}), 400

    provided_order_data = data.get('order_data')
    if provided_order_data is not None and not isinstance(provided_order_data, dict):
        return jsonify({"error": "order_data must be an object"}), 400

    updated_at = datetime.now().isoformat()
    try:
        conn = get_db_connection()
        row = conn.execute(
            'SELECT id, status, order_data, email FROM orders WHERE id = ?',
            (order_id,)
        ).fetchone()

        previous_status = None
        if row:
            previous_status = row['status']
            existing_order_data = {}
            if row['order_data']:
                try:
                    existing_order_data = json.loads(row['order_data'])
                except json.JSONDecodeError:
                    existing_order_data = {}

            final_order_data = provided_order_data if provided_order_data is not None else existing_order_data
            existing_email = row['email'] or final_order_data.get('user_email', '')
            conn.execute(
                'UPDATE orders SET status = ?, updated_at = ?, order_data = ?, email = ? WHERE id = ?',
                (new_status, updated_at, json.dumps(final_order_data), existing_email, order_id)
            )
        else:
            final_order_data = provided_order_data if provided_order_data is not None else {}
            final_email = final_order_data.get('user_email', '')
            conn.execute(
                'INSERT INTO orders (id, status, updated_at, order_data, email) VALUES (?, ?, ?, ?, ?)',
                (order_id, new_status, updated_at, json.dumps(final_order_data), final_email)
            )

        conn.commit()
        conn.close()

        logger.info(f"State Updated for Order {order_id}: {new_status}")

        should_trigger_webhook = (
            new_status == 'SUCCESS'
            and previous_status != 'SUCCESS'
            and bool(final_order_data)
        )
        if should_trigger_webhook:
            logger.info("Payment SUCCESS - Triggering n8n webhook...")
            thread = threading.Thread(target=trigger_n8n_webhook, args=(final_order_data,))
            thread.start()

        return jsonify({
            "message": "Status updated",
            "current_status": new_status,
            "order_id": str(order_id),
            "webhook_triggered": should_trigger_webhook
        })
    except Exception as e:
        logger.error(f"DB Update Error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/<path:path>')
def serve_static(path):
    """Serve static files; registered last so /api/* routes take precedence (clarity for maintainers)."""
    return send_from_directory('.', path)


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    logger.info(f"Starting Payment Polling Server (SQLite) on port {port}...")
    app.run(host='0.0.0.0', port=port)
